# # Import the required libraries
# import pytesseract
# from pdf2image import convert_from_path
# from docx import Document
#
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
#
# # Specify the path of the PDF file
# pdf_file = "Okta_project.pdf"
#
# # Convert the PDF pages to images
# pages = convert_from_path(pdf_file, 300)
#
# # Create a new Word document
# doc = Document()
#
# # Loop over the pages
# for page in pages:
#     # Extract the text from the image using Tesseract OCR
#     text = pytesseract.image_to_string(page, lang="eng")
#     # Add the text to the Word document
#     doc.add_paragraph(text)
#
# # Save the Word document
# doc.save("Okta_project.docx")
#
#
#
# Import the required libraries
import pytesseract
from pdf2image import convert_from_path
from docx import Document

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Specify the path of the PDF file
pdf_file = "Okta_project.pdf"

# Convert the PDF pages to images using the poppler_path argument
poppler_path = r"C:\Program Files\Release-23.07.0-0\poppler-23.07.0\Library\bin"
images = convert_from_path(pdf_file, 500, poppler_path=poppler_path)

# Create a new Word document
doc = Document()

# Loop over the images
for i in range(len(images)):
    # Save the image as a JPEG file
    image_file = "page" + str(i) + ".jpg"
    images[i].save(image_file, "JPEG")
    # Extract the text from the image using Tesseract OCR
    text = pytesseract.image_to_string(image_file, lang="eng")
    # Add the text to the Word document
    doc.add_paragraph(text)

# Save the Word document
doc.save("output.docx")